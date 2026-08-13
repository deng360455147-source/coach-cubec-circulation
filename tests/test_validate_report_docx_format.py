from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from scripts.validate_report_docx_format import validate_docx


def set_style_font(style, east_asia: str, latin: str, size_pt: float, bold: bool) -> None:
    style.font.name = latin
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)


def set_first_line_chars(style, chars_hundredths: int) -> None:
    ppr = style._element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars_hundredths))


def add_seq_field(paragraph, label: str, number: int, title: str) -> None:
    paragraph.add_run(label)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" SEQ {label} \\* ARABIC "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = str(number)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, result, end):
        run._r.append(node)
    paragraph.add_run(f" {title}")


def build_docx(path: Path, compliant: bool = True) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17 if compliant else 2.0)
    section.right_margin = Cm(3.17)

    normal = document.styles["Normal"]
    set_style_font(normal, "宋体" if compliant else "微软雅黑", "Times New Roman", 12, False)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    set_first_line_chars(normal, 200)

    styles = (
        ("Heading 1", "黑体", 14, 6, 6),
        ("Heading 2", "宋体", 12, 6, 0),
        ("Heading 3", "宋体", 12, 3, 0),
    )
    for name, font, size, before, after in styles:
        style = document.styles[name]
        set_style_font(style, font, "Times New Roman", size, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    caption = document.styles["Caption"]
    set_style_font(caption, "黑体", "Times New Roman", 10.5, False)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    document.settings._element.append(update)

    document.add_paragraph("3. 引言", style="Heading 1")
    figure = document.add_paragraph()
    figure.add_run()._r.append(OxmlElement("w:drawing"))
    figure_caption = document.add_paragraph(style="Caption")
    if compliant:
        add_seq_field(figure_caption, "图", 1, "研究框架")
    else:
        figure_caption.add_run("图1 研究框架")

    table_caption = document.add_paragraph(style="Caption")
    add_seq_field(table_caption, "表", 1, "核心指标")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "示例"
    table.cell(1, 1).text = "1"

    document.add_paragraph("11.4 参考文献", style="Heading 2")
    document.add_paragraph("[1] 某机构. 某报告[R]. 2026.")
    document.add_paragraph("11.5 观察/竞品与现场记录", style="Heading 2")
    document.save(path)


class ReportDocxFormatTests(unittest.TestCase):
    def test_compliant_document_passes(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "good.docx"
            build_docx(path, compliant=True)
            self.assertEqual(validate_docx(path, True, True), [])

    def test_wrong_font_margin_and_manual_caption_fail(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "bad.docx"
            build_docx(path, compliant=False)
            errors = validate_docx(path, True, True)
            joined = "\n".join(errors)
            self.assertIn("正文中文字体", joined)
            self.assertIn("left页边距", joined)
            self.assertIn("没有真实SEQ自动编号域", joined)


if __name__ == "__main__":
    unittest.main()
